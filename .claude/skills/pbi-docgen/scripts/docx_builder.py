"""DOCX builder module: assembles branded Word documents from generated prose.

Transforms structured section content into a polished .docx file with cover page,
branded headers/footers, run-level color formatting, XML-shaded tables, and
styled code blocks. Suitable for direct client delivery.

Uses run-level formatting exclusively for colors (D-14/D-15) — never style mutation.
Table header shading uses XML-level OxmlElement (D-16).
Cover page uses different_first_page_header_footer (D-17/D-18).
"""

import os
import re
import sys
from datetime import date

import yaml
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# Relative import for skill scripts
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from scripts.utils import parse_hex_color, ensure_directory


# --- Inline Markdown regex (compiled once at module level for performance) ---
# Priority order: bold+italic (***), bold (**), italic (*), word-boundary italic (_),
# bold (__), bold+italic (___), inline code (backtick).
INLINE_MD_RE = re.compile(
    r'\*\*\*(.+?)\*\*\*'          # ***bold+italic***
    r'|___(.+?)___'                # ___bold+italic___
    r'|\*\*(.+?)\*\*'             # **bold**
    r'|__(.+?)__'                  # __bold__
    r'|\*(.+?)\*'                  # *italic*
    r'|(?:(?<=\s)|(?<=^))_(.+?)_(?=[\s.,;:!?\)]|$)'  # _italic_ (word-boundary aware)
    r'|`([^`]+)`'                  # `code`
)


def _parse_inline_runs(paragraph, text: str, code_style=None) -> None:
    """Parse inline Markdown in *text* and add formatted runs to *paragraph*.

    Handles: ***bold+italic***, **bold**, *italic*, _italic_ (word-boundary),
    ``code``, and plain text segments.

    Args:
        paragraph: A python-docx Paragraph object to append runs to.
        text: The source text possibly containing inline Markdown.
        code_style: Optional Code paragraph style (used for font settings on code spans).
    """
    last_end = 0
    for m in INLINE_MD_RE.finditer(text):
        # Add any plain text before this match
        if m.start() > last_end:
            paragraph.add_run(text[last_end:m.start()])

        # Determine which group matched
        if m.group(1) is not None:
            # ***bold+italic***
            run = paragraph.add_run(m.group(1))
            run.font.bold = True
            run.font.italic = True
        elif m.group(2) is not None:
            # ___bold+italic___
            run = paragraph.add_run(m.group(2))
            run.font.bold = True
            run.font.italic = True
        elif m.group(3) is not None:
            # **bold**
            run = paragraph.add_run(m.group(3))
            run.font.bold = True
        elif m.group(4) is not None:
            # __bold__
            run = paragraph.add_run(m.group(4))
            run.font.bold = True
        elif m.group(5) is not None:
            # *italic*
            run = paragraph.add_run(m.group(5))
            run.font.italic = True
        elif m.group(6) is not None:
            # _italic_ (word-boundary)
            run = paragraph.add_run(m.group(6))
            run.font.italic = True
        elif m.group(7) is not None:
            # `code`
            run = paragraph.add_run(m.group(7))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        last_end = m.end()

    # Add any remaining plain text after the last match
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _add_inline_md_paragraph(doc, line: str, code_style=None):
    """Create a paragraph from *line*, converting inline Markdown to Word runs.

    Handles leading ``- `` as a bullet indicator (uses 'List Bullet' style if
    available, otherwise strips the dash and uses a normal paragraph).

    Args:
        doc: The python-docx Document instance.
        line: A single line of text, possibly with inline Markdown.
        code_style: Optional Code paragraph style for code span formatting.

    Returns:
        The created Paragraph object.
    """
    is_bullet = line.startswith('- ')
    content = line[2:] if is_bullet else line

    if is_bullet:
        # Try to use 'List Bullet' style; fall back to normal paragraph
        try:
            para = doc.add_paragraph(style='List Bullet')
        except KeyError:
            para = doc.add_paragraph()
    else:
        para = doc.add_paragraph()

    _parse_inline_runs(para, content, code_style)
    return para


# --- Language-aware date formatting (D-06, Pattern 3) ---
# Uses month-name dict instead of locale.setlocale() (anti-pattern: platform-dependent, thread-unsafe)
FR_MONTHS = {
    1: 'janvier', 2: 'fevrier', 3: 'mars', 4: 'avril',
    5: 'mai', 6: 'juin', 7: 'juillet', 8: 'aout',
    9: 'septembre', 10: 'octobre', 11: 'novembre', 12: 'decembre',
}


def format_date(d: date, language: str) -> str:
    """Format date per language: FR='le 2 avril 2026', EN='April 02, 2026'."""
    if language == "FR":
        day = '1er' if d.day == 1 else str(d.day)
        return f"le {day} {FR_MONTHS[d.month]} {d.year}"
    return d.strftime('%B %d, %Y')


# --- Cover page boilerplate translations (D-07, Pattern 5) ---
COVER_BOILERPLATE = {
    'EN': {
        'version_prefix': 'Version',
        'prepared_for': 'Prepared for',
        'confidential': 'Confidential',
        'toc_heading': 'Table of Contents',
    },
    'FR': {
        'version_prefix': 'Version',
        'prepared_for': 'Prepare pour',
        'confidential': 'Confidentiel',
        'toc_heading': 'Table des matieres',
    },
}


# --- Section heading map loader and label getter (D-05, Pattern 4) ---
_HEADING_MAP = None


def _load_heading_map():
    global _HEADING_MAP
    if _HEADING_MAP is None:
        map_path = os.path.normpath(
            os.path.join(os.path.dirname(__file__), '..', 'references', 'section_heading_map.yaml')
        )
        with open(map_path, 'r', encoding='utf-8') as f:
            _HEADING_MAP = yaml.safe_load(f)
    return _HEADING_MAP


def get_section_label(section_id: str, language: str) -> str:
    """Get section heading label in the appropriate language."""
    heading_map = _load_heading_map()
    for section in heading_map.get('sections', []):
        if section['id'] == section_id:
            if language == 'FR':
                return section.get('label_fr', section['label'])
            return section['label']
    return section_id.title()


# Ordered list of sections — document renders in this sequence.
# Sections not present in the content dict are silently skipped (D-12).
SECTION_ORDER = ['overview', 'sources', 'dataflows', 'mquery', 'datamodel', 'maintenance']


def _add_colored_heading(doc: Document, text: str, level: int, color_hex: str):
    """Add a heading with brand color applied at run level (D-14/D-15).

    NEVER mutates built-in styles. Creates heading without text argument
    (Pitfall 7), then adds a run and sets run.font.color.rgb directly.

    Args:
        doc: The python-docx Document instance.
        text: Heading text content.
        level: Heading level (1-4).
        color_hex: Hex color string (e.g. '#1B365D' or '1B365D').

    Returns:
        The heading Paragraph object.
    """
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    r, g, b = parse_hex_color(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)
    if level == 1:
        run.font.bold = True
    return heading


def _set_cell_background(cell, color_hex: str) -> None:
    """Set background color on a table cell via XML shading element (D-16).

    MUST create a NEW parse_xml element for EVERY cell — never reuse XML
    elements across cells (Pitfall 3: lxml silently reparents shared elements).

    Args:
        cell: A python-docx table Cell object.
        color_hex: Hex color string (with or without '#').
    """
    color_clean = color_hex.lstrip('#')
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_clean}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_paragraph_shading(paragraph, color_hex: str) -> None:
    """Set background shading on a paragraph via XML (Pattern 5).

    Used for code block backgrounds. Same pattern as cell shading but
    applied to paragraph properties.

    Args:
        paragraph: A python-docx Paragraph object.
        color_hex: Hex color string (with or without '#').
    """
    color_clean = color_hex.lstrip('#')
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_clean}" w:val="clear"/>'
    )
    paragraph._element.get_or_add_pPr().append(shading_elm)


def _add_paragraph_border(paragraph, color_hex: str = "D0D0D0") -> None:
    """Add a thin border around a paragraph (for code block styling).

    Creates pBdr XML with top/left/bottom/right single borders.

    Args:
        paragraph: A python-docx Paragraph object.
        color_hex: Border color hex string without '#'.
    """
    color_clean = color_hex.lstrip('#')
    pPr = paragraph._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="{color_clean}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="4" w:color="{color_clean}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="{color_clean}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="4" w:color="{color_clean}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)


def _create_code_style(document: Document):
    """Create a 'Code' paragraph style for code blocks (D-19).

    Font: Courier New, 9pt, dark grey (#333333).
    Background shading and borders are applied per-paragraph via XML,
    not in the style definition.

    Args:
        document: The python-docx Document instance.

    Returns:
        The Style object for 'Code' paragraphs.
    """
    styles = document.styles
    # Return existing style if already created
    try:
        return styles['Code']
    except KeyError:
        pass
    code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.base_style = styles['Normal']
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return code_style


def _build_cover_page(doc: Document, config: dict, primary_color: str) -> None:
    """Build the cover page with logos, client info, and brand colors (D-17).

    Uses the document's first section. Sets different_first_page_header_footer
    to True (D-18) so body pages get distinct headers/footers.

    Cover page layout:
    1. Top spacing
    2. Client logo (centered, large)
    3. Client name (24pt, bold, primary color, centered)
    4. Report name (18pt, primary color, centered)
    5. Version and date (12pt, grey, centered)
    6. Company logo (right-aligned, bottom area)
    7. Section break for body content

    Args:
        doc: The python-docx Document instance.
        config: Config dict with client_name, client_logo, company_logo, etc.
        primary_color: Primary brand color hex string.
    """
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Top spacing
    for _ in range(3):
        doc.add_paragraph('')

    # Client logo — centered, large
    if config.get('client_logo') and os.path.isfile(config['client_logo']):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = logo_para.add_run()
        run.add_picture(config['client_logo'], width=Inches(3))

    # Spacing after logo
    for _ in range(2):
        doc.add_paragraph('')

    # Client name — 24pt, bold, primary color, centered
    r, g, b = parse_hex_color(primary_color)
    client_para = doc.add_paragraph()
    client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    client_run = client_para.add_run(config.get('client_name', ''))
    client_run.font.size = Pt(24)
    client_run.font.bold = True
    client_run.font.color.rgb = RGBColor(r, g, b)

    # Report name — 18pt, primary color, centered
    report_para = doc.add_paragraph()
    report_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_run = report_para.add_run(config.get('report_name', ''))
    report_run.font.size = Pt(18)
    report_run.font.color.rgb = RGBColor(r, g, b)

    # Spacing
    doc.add_paragraph('')

    # Version and date — 12pt, grey, centered (language-aware per D-06/D-07)
    language = config.get('language', 'EN')
    boilerplate = COVER_BOILERPLATE.get(language, COVER_BOILERPLATE['EN'])
    date_str = format_date(date.today(), language)
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_text = f"{boilerplate['version_prefix']} {config.get('version', '1.0')} | {date_str}"
    version_run = version_para.add_run(version_text)
    version_run.font.size = Pt(12)
    version_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Prepared for line
    prep_para = doc.add_paragraph()
    prep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prep_run = prep_para.add_run(f"{boilerplate['prepared_for']} {config.get('client_name', '')}")
    prep_run.font.size = Pt(12)
    prep_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Push company logo toward bottom with spacing
    for _ in range(8):
        doc.add_paragraph('')

    # Company logo — right-aligned, ~1.5 inches width
    if config.get('company_logo') and os.path.isfile(config['company_logo']):
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = company_para.add_run()
        run.add_picture(config['company_logo'], width=Inches(1.5))

    # Section break for body content
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def _setup_headers_footers(doc: Document, config: dict) -> None:
    """Configure headers and footers for body pages (D-18, Pattern 7, Pitfall 2).

    Body section (sections[1]) gets:
    - Header: company logo left-aligned, ~1 inch width
    - Footer: centered page number field

    Both header and footer are explicitly unlinked from previous section
    (is_linked_to_previous = False) to prevent cover page content leaking.

    Args:
        doc: The python-docx Document instance.
        config: Config dict with company_logo path.
    """
    if len(doc.sections) < 2:
        return

    body_section = doc.sections[1]

    # Header — company logo, left-aligned
    header = body_section.header
    header.is_linked_to_previous = False
    if config.get('company_logo') and os.path.isfile(config['company_logo']):
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = header_para.add_run()
        run.add_picture(config['company_logo'], width=Inches(1.0))

    # Footer — centered page number
    footer = body_section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_simple = parse_xml(
        f'<w:fldSimple {nsdecls("w")} w:instr="PAGE \\* MERGEFORMAT"/>'
    )
    footer_para._element.append(fld_simple)


def _compute_tint(color_hex: str, mix_factor: float = 0.15) -> str:
    """Compute a light tint of a color by mixing with white.

    Args:
        color_hex: Base color hex string (with or without '#').
        mix_factor: How much of the base color to keep (0.15 = 15% color, 85% white).

    Returns:
        Hex color string without '#' prefix (e.g. 'E8EDF3').
    """
    r, g, b = parse_hex_color(color_hex)
    tint_r = int(r * mix_factor + 255 * (1 - mix_factor))
    tint_g = int(g * mix_factor + 255 * (1 - mix_factor))
    tint_b = int(b * mix_factor + 255 * (1 - mix_factor))
    return f"{tint_r:02X}{tint_g:02X}{tint_b:02X}"


def _parse_table_block(lines: list[str]) -> list[list[str]]:
    """Parse pipe-delimited table lines into a list of row-cell lists.

    Handles standard Markdown table format with header, separator, and body rows.
    Strips leading/trailing pipes and whitespace from each cell.

    Args:
        lines: List of pipe-delimited table lines.

    Returns:
        List of rows, each row a list of cell text strings.
        The separator row (containing ---) is excluded.
    """
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Skip separator rows (e.g. |---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', stripped):
            continue
        # Split by pipe, strip whitespace, remove empty leading/trailing from split
        cells = [c.strip() for c in stripped.split('|')]
        # Remove empty strings from leading/trailing pipes
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        if cells:
            rows.append(cells)
    return rows


def _add_table(doc: Document, rows: list[list[str]], primary_color: str, accent_color: str,
               code_style=None) -> None:
    """Add a formatted Word table from parsed row data.

    Header row: white text on primary brand color background.
    Body rows: alternating white and light accent tint (D-16).
    Cell text is parsed for inline Markdown formatting.

    Args:
        doc: The python-docx Document instance.
        rows: List of rows from _parse_table_block.
        primary_color: Primary brand color hex.
        accent_color: Accent brand color hex.
        code_style: Optional Code paragraph style for inline code spans.
    """
    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    tint_hex = _compute_tint(accent_color)

    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = row.cells[col_idx]
                cell.text = ''  # clear default
                para = cell.paragraphs[0]
                _parse_inline_runs(para, cell_text, code_style)
                if row_idx == 0:
                    # Header row — white text on primary color background
                    _set_cell_background(cell, primary_color)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.bold = True
                elif row_idx % 2 == 0:
                    # Even body rows (0-indexed, row_idx 2, 4, ...) — light accent tint
                    _set_cell_background(cell, tint_hex)


def _parse_and_add_prose(doc: Document, prose_text: str, primary_color: str,
                         accent_color: str, code_style) -> None:
    """Parse generated prose and add content to the document.

    Detects and handles three content types:
    1. TABLE: markers with pipe-delimited rows -> Word tables
    2. CODE_BLOCK: markers or triple-backtick fences -> Code-styled paragraphs
    3. Regular text -> Normal paragraphs

    Sub-headings (## or ###) within prose are rendered as level 2 or 3 headings
    with brand color.

    Args:
        doc: The python-docx Document instance.
        prose_text: The raw prose text from content generation.
        primary_color: Primary brand color hex.
        accent_color: Accent brand color hex.
        code_style: The Code paragraph style object.
    """
    lines = prose_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip fully empty lines
        if not line.strip():
            i += 1
            continue

        # TABLE: marker — collect until end marker or empty line
        if line.strip().startswith('TABLE:'):
            table_lines = []
            i += 1
            while i < len(lines):
                tl = lines[i]
                if tl.strip().startswith('END_TABLE') or (not tl.strip() and table_lines):
                    # Check if next line is also empty or non-table
                    if not tl.strip():
                        # Peek ahead: if next line has pipes, keep going
                        if i + 1 < len(lines) and '|' in lines[i + 1]:
                            i += 1
                            continue
                    i += 1
                    break
                if '|' in tl:
                    table_lines.append(tl)
                i += 1
            rows = _parse_table_block(table_lines)
            _add_table(doc, rows, primary_color, accent_color, code_style)
            continue

        # Pipe-delimited table without explicit TABLE: marker
        if '|' in line and re.match(r'^\s*\|', line):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                rows = _parse_table_block(table_lines)
                _add_table(doc, rows, primary_color, accent_color, code_style)
            else:
                # Not a real table, add as normal text
                for tl in table_lines:
                    _add_inline_md_paragraph(doc, tl.strip(), code_style)
            continue

        # CODE_BLOCK: marker
        if line.strip().startswith('CODE_BLOCK:'):
            i += 1
            while i < len(lines):
                cl = lines[i]
                if cl.strip().startswith('END_CODE_BLOCK') or (not cl.strip()):
                    i += 1
                    break
                para = doc.add_paragraph(cl, style=code_style)
                _set_paragraph_shading(para, 'F2F2F2')
                _add_paragraph_border(para)
                i += 1
            continue

        # Triple-backtick fenced code block
        if line.strip().startswith('```'):
            i += 1
            while i < len(lines):
                cl = lines[i]
                if cl.strip().startswith('```'):
                    i += 1
                    break
                para = doc.add_paragraph(cl, style=code_style)
                _set_paragraph_shading(para, 'F2F2F2')
                _add_paragraph_border(para)
                i += 1
            continue

        # Sub-headings within prose (### level 3, ## level 2)
        heading_match = re.match(r'^(#{2,3})\s+(.+)$', line.strip())
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            _add_colored_heading(doc, heading_text, level, primary_color)
            i += 1
            continue

        # Regular text paragraph (inline Markdown -> Word runs)
        _add_inline_md_paragraph(doc, line.strip(), code_style)
        i += 1


def _insert_toc_field(doc):
    """Insert a Table of Contents field into the document body (D-01/D-02/D-03).

    Creates a Structured Document Tag (SDT) wrapping a TOC field instruction:
    TOC \\o "1-3" \\h \\z \\u

    Word displays 'Update this field' on open; user presses F9 to refresh.
    The \\o "1-3" switch collects Heading 1-3 entries.
    The \\h switch creates hyperlinks, \\z hides tab leaders in web view,
    \\u uses applied paragraph outline levels.
    """
    sdt = OxmlElement('w:sdt')
    sdtpr = OxmlElement('w:sdtPr')
    docpartobj = OxmlElement('w:docPartObj')
    docpartgallery = OxmlElement('w:docPartGallery')
    docpartgallery.set(qn('w:val'), 'Table of Contents')
    docpartunique = OxmlElement('w:docPartUnique')
    docpartunique.set(qn('w:val'), 'true')
    docpartobj.append(docpartgallery)
    docpartobj.append(docpartunique)
    sdtpr.append(docpartobj)
    sdt.append(sdtpr)

    sdtcontent = OxmlElement('w:sdtContent')

    # TOC field paragraph: begin -> instrText -> separate -> end
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')

    fldchar_begin = OxmlElement('w:fldChar')
    fldchar_begin.set(qn('w:fldCharType'), 'begin')

    instrtext = OxmlElement('w:instrText')
    instrtext.set(qn('xml:space'), 'preserve')
    instrtext.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldchar_separate = OxmlElement('w:fldChar')
    fldchar_separate.set(qn('w:fldCharType'), 'separate')

    fldchar_end = OxmlElement('w:fldChar')
    fldchar_end.set(qn('w:fldCharType'), 'end')

    r.append(fldchar_begin)
    r.append(instrtext)
    r.append(fldchar_separate)
    r.append(fldchar_end)
    p.append(r)
    sdtcontent.append(p)
    sdt.append(sdtcontent)

    # Append SDT to document body
    doc.element.body.append(sdt)


def build_docx(sections: dict, config: dict) -> str:
    """Build a complete branded DOCX document from generated section content.

    Assembles cover page, headers/footers, and all content sections into a
    polished Word document suitable for direct client delivery.

    Args:
        sections: Dict of section content from content_generator.generate_all_sections.
            Keys are section IDs (e.g. 'overview', 'sources'), values are dicts
            with 'label' (display heading) and 'prose' (generated text).
        config: JSON config dict with keys per D-22:
            - client_name: str
            - client_logo: str (path)
            - company_logo: str (path)
            - primary_color: str (hex, e.g. '#1B365D')
            - accent_color: str (hex, e.g. '#4A90D9')
            - report_name: str
            - version: str
            - output_dir: str

    Returns:
        Absolute path to the saved .docx file.
    """
    doc = Document()
    code_style = _create_code_style(doc)

    primary_color = config.get('primary_color', '#1B365D')
    accent_color = config.get('accent_color', '#4A90D9')

    # Build cover page (first section)
    _build_cover_page(doc, config, primary_color)

    # Language-aware TOC and section headings (D-05/D-07)
    language = config.get('language', 'EN')
    boilerplate = COVER_BOILERPLATE.get(language, COVER_BOILERPLATE['EN'])

    # NEW: TOC page (D-01/D-02/D-03)
    _add_colored_heading(doc, boilerplate['toc_heading'], 1, primary_color)
    _insert_toc_field(doc)
    doc.add_page_break()

    # Set up headers and footers for body section
    _setup_headers_footers(doc, config)

    # Add content sections in defined order
    for section_id in SECTION_ORDER:
        if section_id not in sections:
            continue
        section_data = sections[section_id]
        # Override label with language-appropriate heading (D-05, Pitfall 3)
        label = get_section_label(section_id, language)
        prose = section_data.get('prose', '')

        # Section heading — level 1, primary color
        _add_colored_heading(doc, label, 1, primary_color)

        # Section content
        if prose:
            _parse_and_add_prose(doc, prose, primary_color, accent_color, code_style)

    # Construct output filename: {client}_{report}_v{version}_{date}.docx
    client_name = config.get('client_name', 'Client').replace(' ', '_')
    report_name = config.get('report_name', 'Report').replace(' ', '_')
    version = config.get('version', '1.0')
    today = date.today().strftime('%Y-%m-%d')
    filename = f"{client_name}_{report_name}_v{version}_{today}.docx"

    # Ensure output directory exists
    output_dir = config.get('output_dir', 'docsgen-output')
    ensure_directory(output_dir)

    output_path = os.path.abspath(os.path.join(output_dir, filename))
    doc.save(output_path)

    print(f"Document saved: {output_path} ({os.path.getsize(output_path)} bytes)",
          file=sys.stderr)

    return output_path
